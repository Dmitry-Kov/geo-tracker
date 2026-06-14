# -*- coding: utf-8 -*-
"""
GEO tracker report export to a branded Optimize.uz DOCX.

Usage:
  python report_docx.py                 → geo_report_en_DATE.docx
  python report_docx.py --lang ru       → geo_report_ru_DATE.docx (for a Russian-reading boss)
  python report_docx.py --out path

The report chrome is bilingual (--lang en|ru); niche titles always come
from queries.py. The style mirrors dashboard.html: ink #16181D, accent
#FF5C1F, Space Grotesk headings, Inter body text, JetBrains Mono figures.
"""

import argparse
import csv
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from queries import DOMAINS, NICHES

BASE_DIR = Path(__file__).resolve().parent
RESULTS_CSV = BASE_DIR / "results.csv"

INK = RGBColor(0x16, 0x18, 0x1D)
ACCENT = RGBColor(0xFF, 0x5C, 0x1F)
MUTED = RGBColor(0x6B, 0x71, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT_SOFT = "FFF0E9"   # w:shd fills take a plain hex string
LINE = "E4E6E0"

F_HEAD = "Space Grotesk"
F_BODY = "Inter"
F_MONO = "JetBrains Mono"


# ---------------------------------------------------------------------------
# Data (same logic as report_html)
# ---------------------------------------------------------------------------

def _hit(row, domain):
    return row.get(f"src_{domain}") == "1" or row.get(f"txt_{domain}") == "1"


def load_ok_rows():
    with RESULTS_CSV.open(newline="", encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))
    return rows, [r for r in rows if r["status"] == "ok"]


def _ts(row):
    return row.get("run_ts") or row.get("run_date") or ""


def _latest_per_pair(ok):
    """Freshest ok answer per (engine, niche, query) — the current-visibility
    snapshot, so a partial gap-fill day doesn't shrink the report to that day."""
    latest = {}
    for r in ok:
        key = (r["engine"], r["niche"], r["query"])
        if key not in latest or _ts(r) >= _ts(latest[key]):
            latest[key] = r
    return list(latest.values())


def sov(rows, *, niche=None, engine=None):
    subset = [r for r in rows
              if (niche is None or r["niche"] == niche)
              and (engine is None or r["engine"] == engine)]
    if not subset:
        return None
    return {d: round(100 * sum(_hit(r, d) for r in subset) / len(subset), 1)
            for d in DOMAINS}


# ---------------------------------------------------------------------------
# Low-level formatting helpers
# ---------------------------------------------------------------------------

def shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def style_run(run, *, font=F_BODY, size=10, bold=False, color=INK):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def para(doc, text="", *, font=F_BODY, size=10, bold=False, color=INK,
         space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        style_run(p.add_run(text), font=font, size=size, bold=bold, color=color)
    return p


def heading(doc, text, *, size=13):
    return para(doc, text, font=F_HEAD, size=size, bold=True,
                space_before=14, space_after=4)


def hairline(paragraph, color=LINE):
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pbdr)


def make_table(doc, n_rows, n_cols):
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # repaint the Table Grid borders in the brand hairline color
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), LINE)
        borders.append(el)
    tbl_pr.append(borders)
    return table


def fill_cell(cell, text, *, font=F_BODY, size=9, bold=False, color=INK,
              fill=None, center=True):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(text), font=font, size=size, bold=bold, color=color)
    if fill:
        shade(cell, fill)


def header_row(table, labels):
    for i, label in enumerate(labels):
        fill_cell(table.rows[0].cells[i], label,
                  font=F_MONO, size=8, bold=True, color=WHITE, fill="16181D")


def pct_cell(cell, value, *, primary=False):
    """Percentage cell: fill intensity by visibility level, marker on the primary."""
    if value >= 50:
        fill, color, bold = "FF5C1F", WHITE, True
    elif value >= 20:
        fill, color, bold = "FFC9B0", INK, True
    elif value > 0:
        fill, color, bold = ACCENT_SOFT, INK, True
    else:
        fill, color, bold = None, MUTED, False
    text = f"{value:.0f}%" + (" ★" if primary else "")
    fill_cell(cell, text, font=F_MONO, size=9, bold=bold, color=color, fill=fill)


# ---------------------------------------------------------------------------
# Localized strings. Niche titles come from queries.py (the user's config),
# so only the report chrome is localized here. {…} are str.format placeholders.
# ---------------------------------------------------------------------------

STRINGS = {
    "en": {
        "header_suffix": "  GEO Monitor",
        "subtitle": "Domain citation report in AI answers · current snapshot through {latest} · generated {ts}",
        "kpi_answers": "pairs in snapshot",
        "kpi_engines": "engines: {engines}",
        "kpi_hits": "domain hits",
        "kpi_dates": "dates in history ({first} … {latest})",
        "h_visibility": "Visibility by niche",
        "visibility_note": "Share of the niche's queries where the domain appeared in the "
                           "answer's sources (src) or text (txt). ★ — niche's primary domain.",
        "col_niche": "niche",
        "h_leaderboard": "Domain leaderboard",
        "lb_cols": ["domain", "SOV", "note"],
        "lb_niches": "niches: {names}",
        "lb_not_cited": "not cited",
        "h_engines": "Engine comparison",
        "ec_cols": ["engine", "n"],
        "h_observations": "Key observations",
        "obs_consensus": "{d} in niche “{title}” — cited by all engines",
        "obs_single": "{d} in niche “{title}” — only {who}",
        "obs_zeros": "✗ Not cited by any engine: {domains}",
        "h_hits": "Queries with hits (current snapshot)",
        "hits_cols": ["engine", "niche", "query", "domains"],
        "hits_none": "No hits.",
        "footer": "Optimize.uz · GEO tracker · official APIs only (Gemini, Perplexity, "
                  "OpenAI). AI is non-deterministic: read the trend across several runs, "
                  "not a single measurement.",
    },
    "ru": {
        "header_suffix": "  GEO-монитор",
        "subtitle": "Отчёт о цитируемости доменов в ответах ИИ · текущий срез по {latest} · "
                    "сформирован {ts}",
        "kpi_answers": "пар в срезе",
        "kpi_engines": "движков: {engines}",
        "kpi_hits": "попаданий доменов",
        "kpi_dates": "дат в истории ({first} … {latest})",
        "h_visibility": "Видимость по нишам",
        "visibility_note": "Доля запросов ниши, где домен попал в источники (src) или текст "
                           "ответа (txt). ★ — целевой домен ниши.",
        "col_niche": "ниша",
        "h_leaderboard": "Общий зачёт по доменам",
        "lb_cols": ["домен", "SOV", "комментарий"],
        "lb_niches": "ниши: {names}",
        "lb_not_cited": "не цитируется",
        "h_engines": "Сравнение движков",
        "ec_cols": ["движок", "n"],
        "h_observations": "Ключевые наблюдения",
        "obs_consensus": "{d} в нише «{title}» — цитируют все движки",
        "obs_single": "{d} в нише «{title}» — только {who}",
        "obs_zeros": "✗ Не цитируются ни одним движком: {domains}",
        "h_hits": "Запросы с попаданиями (текущий срез)",
        "hits_cols": ["движок", "ниша", "запрос", "домены"],
        "hits_none": "Попаданий нет.",
        "footer": "Optimize.uz · GEO-трекер · только официальные API (Gemini, Perplexity, "
                  "OpenAI). ИИ недетерминирован: смотрите тренд по нескольким прогонам, "
                  "а не единичный замер.",
    },
}


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def build_docx(out_path: Path, lang: str = "en") -> Path:
    T = STRINGS[lang]
    all_rows, ok = load_ok_rows()
    dates = sorted({r["run_date"] for r in all_rows})
    latest = dates[-1]
    # Current snapshot: freshest answer per pair across all history, so the
    # report reflects the whole matrix, not just the last (maybe partial) day.
    current = _latest_per_pair(ok)
    engines = sorted({r["engine"] for r in current})

    doc = Document()
    for section in doc.sections:
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.left_margin = section.right_margin = Cm(2)
        section.top_margin, section.bottom_margin = Cm(1.8), Cm(1.8)

    # Header: Optimize<.>uz GEO monitor, same as the dashboard
    p = doc.add_paragraph()
    style_run(p.add_run("Optimize"), font=F_HEAD, size=22, bold=True)
    style_run(p.add_run("."), font=F_HEAD, size=22, bold=True, color=ACCENT)
    style_run(p.add_run("uz"), font=F_HEAD, size=22, bold=True)
    style_run(p.add_run(T["header_suffix"]), font=F_HEAD, size=22, bold=True)
    hairline(p, color="16181D")
    para(doc,
         T["subtitle"].format(latest=latest,
                              ts=datetime.now().strftime('%d.%m.%Y %H:%M')),
         font=F_MONO, size=8.5, color=MUTED, space_after=12)

    # KPI row
    kpi = make_table(doc, 2, 4)
    for i, (val, label) in enumerate([
            (str(len(current)), T["kpi_answers"].format(latest=latest)),
            (str(len(engines)), T["kpi_engines"].format(engines=", ".join(engines))),
            (str(sum(1 for r in current for d in DOMAINS if _hit(r, d))),
             T["kpi_hits"]),
            (f"{len(dates)}", T["kpi_dates"].format(first=dates[0], latest=latest))]):
        fill_cell(kpi.rows[0].cells[i], val, font=F_MONO, size=16, bold=True,
                  color=ACCENT if i == 2 else INK)
        fill_cell(kpi.rows[1].cells[i], label, size=8, color=MUTED)

    # Visibility matrix by niche
    heading(doc, T["h_visibility"])
    para(doc, T["visibility_note"], size=9, color=MUTED)
    short = {d: d.split(".")[0] for d in DOMAINS}
    m = make_table(doc, len([n for n in NICHES if sov(current, niche=n)]) + 1,
                   len(DOMAINS) + 1)
    header_row(m, [T["col_niche"]] + [short[d] for d in DOMAINS])
    r_i = 1
    for niche, cfg in NICHES.items():
        by_dom = sov(current, niche=niche)
        if not by_dom:
            continue
        fill_cell(m.rows[r_i].cells[0], cfg["title"], size=9, bold=True, center=False)
        for c_i, d in enumerate(DOMAINS, start=1):
            pct_cell(m.rows[r_i].cells[c_i], by_dom[d],
                     primary=d == cfg["primary_domain"])
        r_i += 1

    # Overall leaderboard
    heading(doc, T["h_leaderboard"])
    total = sov(current) or {}
    lb = sorted(total.items(), key=lambda x: -x[1])
    t = make_table(doc, len(lb) + 1, 3)
    header_row(t, T["lb_cols"])
    for i, (d, pct) in enumerate(lb, start=1):
        niches_hit = sorted({r["niche"] for r in current if _hit(r, d)})
        note = (T["lb_niches"].format(names=", ".join(NICHES[n]["title"] for n in niches_hit))
                if niches_hit else T["lb_not_cited"])
        fill_cell(t.rows[i].cells[0], d, font=F_MONO, size=9, center=False)
        pct_cell(t.rows[i].cells[1], pct)
        fill_cell(t.rows[i].cells[2], note, size=8.5,
                  color=INK if niches_hit else MUTED, center=False)

    # Per-engine breakdown
    heading(doc, T["h_engines"])
    e_rows = [e for e in engines if sov(current, engine=e)]
    et = make_table(doc, len(e_rows) + 1, len(DOMAINS) + 2)
    header_row(et, T["ec_cols"] + [short[d] for d in DOMAINS])
    for i, e in enumerate(e_rows, start=1):
        sub_n = sum(1 for r in current if r["engine"] == e)
        by_dom = sov(current, engine=e)
        fill_cell(et.rows[i].cells[0], e, font=F_MONO, size=9, bold=True, center=False)
        fill_cell(et.rows[i].cells[1], str(sub_n), font=F_MONO, size=9, color=MUTED)
        for c_i, d in enumerate(DOMAINS, start=2):
            pct_cell(et.rows[i].cells[c_i], by_dom[d])

    # Auto-derived findings from the data
    heading(doc, T["h_observations"])
    consensus = []
    single = []
    for d in DOMAINS:
        for niche in NICHES:
            who = {r["engine"] for r in current
                   if r["niche"] == niche and _hit(r, d)}
            if len(who) == len(engines) > 1:
                consensus.append(T["obs_consensus"].format(d=d, title=NICHES[niche]["title"]))
            elif len(who) == 1:
                single.append(T["obs_single"].format(
                    d=d, title=NICHES[niche]["title"], who=next(iter(who))))
    zeros = [d for d in DOMAINS if not any(_hit(r, d) for r in current)]
    for line in consensus:
        para(doc, "✓ " + line, size=9.5, space_after=3)
    for line in single:
        para(doc, "◐ " + line, size=9.5, color=MUTED, space_after=3)
    if zeros:
        para(doc, T["obs_zeros"].format(domains=", ".join(zeros)),
             size=9.5, bold=True, color=ACCENT, space_after=3)

    # Hits from the latest run
    heading(doc, T["h_hits"])
    hits = [(r, [d for d in DOMAINS if _hit(r, d)])
            for r in current if any(_hit(r, d) for d in DOMAINS)]
    if hits:
        ht = make_table(doc, len(hits) + 1, 4)
        header_row(ht, T["hits_cols"])
        for i, (r, ds) in enumerate(hits, start=1):
            fill_cell(ht.rows[i].cells[0], r["engine"], font=F_MONO, size=8.5,
                      center=False)
            fill_cell(ht.rows[i].cells[1], r["niche"], size=8.5, center=False)
            fill_cell(ht.rows[i].cells[2], r["query"], size=8.5, center=False)
            fill_cell(ht.rows[i].cells[3], ", ".join(ds), font=F_MONO, size=8.5,
                      color=ACCENT, center=False, fill=ACCENT_SOFT)
    else:
        para(doc, T["hits_none"], size=9.5, color=MUTED)

    # Footer
    p = para(doc, "", space_before=16)
    hairline(p)
    para(doc, T["footer"], font=F_MONO, size=8, color=MUTED)

    doc.save(out_path)
    return out_path


def main():
    parser = argparse.ArgumentParser(description="Export the report to an Optimize.uz DOCX")
    parser.add_argument("--out", help="path for the output .docx")
    parser.add_argument("--lang", choices=sorted(STRINGS), default="en",
                        help="report-chrome language (default: en); niche titles come from queries.py")
    args = parser.parse_args()
    if not RESULTS_CSV.exists():
        raise SystemExit("results.csv not found — run first: python geo_tracker.py run")
    default = BASE_DIR / f"geo_report_{args.lang}_{datetime.now().strftime('%Y-%m-%d')}.docx"
    out = build_docx(Path(args.out) if args.out else default, lang=args.lang)
    print(f"Done: {out}")


if __name__ == "__main__":
    main()
